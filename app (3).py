import streamlit as st
from io import BytesIO
from typing import Tuple
import os

# Third-party imports guarded in a try/except so error messages are helpful
try:
    from PyPDF2 import PdfReader
except Exception as e:
    PdfReader = None
try:
    import docx  # python-docx
except Exception as e:
    docx = None

st.set_page_config(page_title="Advanced File Upload Dashboard", layout="wide")

st.title("Advanced File Upload Dashboard — Read / Write / Append (PDF & DOCX only)")
st.markdown(
    """
This dashboard accepts **PDF** and **Word (.docx)** files only.  
- **Read**: extract and preview text.  
- **Write**: replace file content with your input (download as `.txt` or `.docx`).  
- **Append**: add text to end of file content (download as `.txt` or `.docx`).
"""
)

# --- Helpers ---
MAX_FILE_SIZE_MB = 10

def get_file_meta(uploaded) -> Tuple[str, int]:
    name = uploaded.name
    size = uploaded.size
    return name, size

def allowed_filetype(filename: str) -> bool:
    allowed = (".pdf", ".docx")
    return filename.lower().endswith(allowed)

def extract_text_from_pdf_bytes(data: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("PyPDF2 is not installed. See requirements.")
    try:
        reader = PdfReader(BytesIO(data))
        texts = []
        for p in reader.pages:
            txt = p.extract_text()
            if txt:
                texts.append(txt)
        return "\n\n".join(texts).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {e}")

def extract_text_from_docx_bytes(data: bytes) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed. See requirements.")
    try:
        bio = BytesIO(data)
        document = docx.Document(bio)
        paragraphs = [p.text for p in document.paragraphs if p.text]
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")

def create_docx_bytes_from_text(text: str) -> bytes:
    if docx is None:
        raise RuntimeError("python-docx is not installed. See requirements.")
    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.read()

# --- UI ---
uploaded_file = st.file_uploader(
    "Upload your file (PDF or DOCX)",
    type=["pdf", "docx"],
    accept_multiple_files=False
)

operation = st.selectbox("Select Operation", ("Read", "Write", "Append"))

user_text = st.text_area("Enter text (Used for Write & Append)", height=200)

# Extra options
st.sidebar.header("Advanced options")
download_as_docx = st.sidebar.checkbox("Offer download as .docx (when available)", value=True)
download_as_txt = st.sidebar.checkbox("Offer download as .txt", value=True)
show_full_preview = st.sidebar.checkbox("Show full preview (may be long)", value=False)

if st.button("Process"):
    if uploaded_file is None:
        st.error("No file uploaded. Please upload a PDF or DOCX file.")
    else:
        filename, filesize = get_file_meta(uploaded_file)
        if filesize > MAX_FILE_SIZE_MB * 1024 * 1024:
            st.error(f"File is too large ({filesize/1024/1024:.2f} MB). Max allowed is {MAX_FILE_SIZE_MB} MB.")
        elif not allowed_filetype(filename):
            st.error("Invalid file type. Only .pdf and .docx files are accepted.")
        else:
            raw = uploaded_file.read()
            try:
                if filename.lower().endswith(".pdf"):
                    extracted = extract_text_from_pdf_bytes(raw)
                else:
                    extracted = extract_text_from_docx_bytes(raw)
            except RuntimeError as e:
                st.error(str(e))
                extracted = None
            except Exception as e:
                st.error(f"Unexpected error while extracting text: {e}")
                extracted = None

            if extracted is None:
                st.stop()

            # Perform the operation
            if operation == "Read":
                st.subheader("Extracted content (preview)")
                if not extracted.strip():
                    st.warning("No extractable text found in the uploaded file.")
                if show_full_preview:
                    st.text_area("Full content", extracted, height=400)
                else:
                    st.code("\n".join(extracted.splitlines()[:30]) + ("\n\n... (truncated)" if len(extracted.splitlines())>30 else ""))

                # Offer downloads of raw text
                if download_as_txt:
                    st.download_button(
                        "Download extracted text (.txt)",
                        data=extracted.encode("utf-8"),
                        file_name=f"{os.path.splitext(filename)[0]}_extracted.txt",
                        mime="text/plain"
                    )
                if download_as_docx and docx is not None:
                    try:
                        docx_bytes = create_docx_bytes_from_text(extracted)
                        st.download_button(
                            "Download extracted content (.docx)",
                            data=docx_bytes,
                            file_name=f"{os.path.splitext(filename)[0]}_extracted.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Could not create .docx for download: {e}")

            elif operation == "Write":
                if not user_text.strip():
                    st.error("Write operation requires non-empty text in the input box. Please enter the content you want to save.")
                else:
                    updated = user_text.strip()
                    st.success("Write operation prepared — content will replace file contents for download (original upload is not modified).")
                    st.markdown("**Preview (first 2000 characters):**")
                    st.text(updated[:2000] + ("\n\n... (truncated)" if len(updated)>2000 else ""))

                    if download_as_txt:
                        st.download_button(
                            "Download replaced content (.txt)",
                            data=updated.encode("utf-8"),
                            file_name=f"{os.path.splitext(filename)[0]}_written.txt",
                            mime="text/plain"
                        )
                    if download_as_docx and docx is not None:
                        try:
                            docx_bytes = create_docx_bytes_from_text(updated)
                            st.download_button(
                                "Download replaced content (.docx)",
                                data=docx_bytes,
                                file_name=f"{os.path.splitext(filename)[0]}_written.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"Could not create .docx for download: {e}")

            elif operation == "Append":
                if not user_text.strip():
                    st.error("Append operation requires non-empty text to append.")
                else:
                    appended = extracted + "\n\n" + user_text.strip()
                    st.success("Append operation prepared — combined content available for download.")
                    st.markdown("**Preview (first 2000 characters):**")
                    st.text(appended[:2000] + ("\n\n... (truncated)" if len(appended)>2000 else ""))

                    if download_as_txt:
                        st.download_button(
                            "Download appended content (.txt)",
                            data=appended.encode("utf-8"),
                            file_name=f"{os.path.splitext(filename)[0]}_appended.txt",
                            mime="text/plain"
                        )
                    if download_as_docx and docx is not None:
                        try:
                            docx_bytes = create_docx_bytes_from_text(appended)
                            st.download_button(
                                "Download appended content (.docx)",
                                data=docx_bytes,
                                file_name=f"{os.path.splitext(filename)[0]}_appended.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"Could not create .docx for download: {e}")
