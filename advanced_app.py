import streamlit as st
from io import BytesIO
import os

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

st.set_page_config(page_title="Advanced File Upload Dashboard", layout="wide")
st.title("Advanced File Upload Dashboard — PDF & Word (.docx)")

st.markdown("""
- Upload **PDF** or **Word (.docx)** files.  
- For PDFs: Read-only + convert to Word.  
- For Word: Read, Write, Append available.
""")

# --- Helpers ---
def allowed_filetype(filename):
    return filename.lower().endswith((".pdf", ".docx"))

def extract_pdf_text(data):
    if PdfReader is None:
        st.error("PyPDF2 is not installed.")
        return ""
    reader = PdfReader(BytesIO(data))
    texts = [p.extract_text() or "" for p in reader.pages]
    return "\n\n".join(texts).strip()

def extract_docx_text(data):
    if docx is None:
        st.error("python-docx is not installed.")
        return ""
    document = docx.Document(BytesIO(data))
    return "\n\n".join([p.text for p in document.paragraphs if p.text]).strip()

def create_docx_bytes_from_text(text):
    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.read()

# --- File upload ---
uploaded_file = st.file_uploader("Upload your file (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    filename = uploaded_file.name
    raw = uploaded_file.read()
    
    # Determine file type
    is_pdf = filename.lower().endswith(".pdf")
    is_docx = filename.lower().endswith(".docx")
    
    if not allowed_filetype(filename):
        st.error("Only PDF and Word files are accepted.")
    else:
        # --- PDF file ---
        if is_pdf:
            st.info("PDF detected: Only Read and Convert to Word are available.")
            operation = st.selectbox("Select Operation", ("Read",))
            if operation == "Read":
                text = extract_pdf_text(raw)
                if text:
                    st.subheader("PDF Content Preview")
                    st.text_area("Preview", text[:2000] + ("\n\n... (truncated)" if len(text)>2000 else ""), height=400)
                    # Convert to Word option
                    if st.button("Convert PDF to Word"):
                        docx_bytes = create_docx_bytes_from_text(text)
                        st.download_button(
                            "Download Word File",
                            data=docx_bytes,
                            file_name=f"{os.path.splitext(filename)[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.warning("No extractable text found in PDF.")
        
        # --- DOCX file ---
        elif is_docx:
            st.info("Word file detected: Read, Write, Append are available.")
            operation = st.selectbox("Select Operation", ("Read", "Write", "Append"))
            text = extract_docx_text(raw)
            
            user_text = st.text_area("Enter text (for Write / Append)", height=200)
            
            if st.button("Process"):
                if operation == "Read":
                    st.subheader("Word Content Preview")
                    st.text_area("Preview", text[:2000] + ("\n\n... (truncated)" if len(text)>2000 else ""), height=400)
                elif operation == "Write":
                    if not user_text.strip():
                        st.error("Write operation requires non-empty text.")
                    else:
                        new_content = user_text.strip()
                        st.success("Content ready to save (Write operation).")
                        st.download_button(
                            "Download Written Word File",
                            data=create_docx_bytes_from_text(new_content),
                            file_name=f"{os.path.splitext(filename)[0]}_written.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                elif operation == "Append":
                    if not user_text.strip():
                        st.error("Append operation requires non-empty text.")
                    else:
                        new_content = text + "\n\n" + user_text.strip()
                        st.success("Content ready to save (Append operation).")
                        st.download_button(
                            "Download Appended Word File",
                            data=create_docx_bytes_from_text(new_content),
                            file_name=f"{os.path.splitext(filename)[0]}_appended.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
