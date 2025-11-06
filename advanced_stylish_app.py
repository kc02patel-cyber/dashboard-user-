import streamlit as st
from io import BytesIO
import os

# Check for necessary libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

# --- Page config ---
st.set_page_config(page_title="📄 File Dashboard", page_icon="📂", layout="wide")

# --- Custom CSS for colors and styling ---
st.markdown(
    """
    <style>
    .stApp {
        background-color: #f0f8ff;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        font-size: 16px;
        padding: 10px 24px;
        border-radius: 10px;
    }
    .stDownloadButton>button {
        background-color: #ff8c00;
        color: white;
        font-size: 16px;
        padding: 10px 24px;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True
)

st.title("📂 Advanced File Dashboard")
st.markdown(
    "Upload **PDF** or **Word (.docx)** files and perform operations with ease! 🎉"
)

# --- Helper functions ---
def allowed_filetype(filename):
    return filename.lower().endswith((".pdf", ".docx"))

def extract_pdf_text(data):
    if PdfReader is None:
        st.error("PyPDF2 not installed.")
        return ""
    reader = PdfReader(BytesIO(data))
    texts = [p.extract_text() or "" for p in reader.pages]
    return "\n\n".join(texts).strip()

def extract_docx_text(data):
    if docx is None:
        st.error("python-docx not installed.")
        return ""
    document = docx.Document(BytesIO(data))
    return "\n\n".join([p.text for p in document.paragraphs if p.text]).strip()

def create_docx_bytes_from_text(text):
    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.read()

# --- Sidebar options ---
st.sidebar.header("💡 Options")
show_preview = st.sidebar.checkbox("Show Full Preview", value=False)
download_txt = st.sidebar.checkbox("Offer download as .txt", value=True)
download_docx = st.sidebar.checkbox("Offer download as .docx", value=True)

# --- File upload ---
uploaded_file = st.file_uploader("Upload your file (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    filename = uploaded_file.name
    raw = uploaded_file.read()
    
    is_pdf = filename.lower().endswith(".pdf")
    is_docx = filename.lower().endswith(".docx")
    
    if not allowed_filetype(filename):
        st.error("Only PDF and Word files are accepted.")
    else:
        # --- PDF Section ---
        if is_pdf:
            st.info("PDF detected: Only Read and Convert to Word are available 📖")
            text = extract_pdf_text(raw)
            st.subheader("📄 PDF Content Preview")
            if text:
                if show_preview:
                    st.text_area("Preview", text, height=400)
                else:
                    st.text_area("Preview", text[:2000] + ("\n\n... (truncated)" if len(text)>2000 else ""), height=400)
                # Convert PDF to Word
                if st.button("📝 Convert PDF to Word"):
                    docx_bytes = create_docx_bytes_from_text(text)
                    st.download_button(
                        "💾 Download Word File",
                        data=docx_bytes,
                        file_name=f"{os.path.splitext(filename)[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("No extractable text found in PDF.")

        # --- DOCX Section ---
        elif is_docx:
            st.success("Word file detected: Read / Write / Append enabled ✨")
            operation = st.selectbox("Select Operation", ("Read", "Write", "Append"))
            text = extract_docx_text(raw)
            
            user_text = st.text_area("Enter text for Write / Append", height=200)
            
            if st.button("🚀 Process"):
                if operation == "Read":
                    st.subheader("📄 Word Content Preview")
                    if show_preview:
                        st.text_area("Preview", text, height=400)
                    else:
                        st.text_area("Preview", text[:2000] + ("\n\n... (truncated)" if len(text)>2000 else ""), height=400)
                    # Download options
                    if download_txt:
                        st.download_button(
                            "💾 Download as .txt",
                            data=text.encode("utf-8"),
                            file_name=f"{os.path.splitext(filename)[0]}_read.txt",
                            mime="text/plain"
                        )
                    if download_docx and docx is not None:
                        st.download_button(
                            "💾 Download as .docx",
                            data=create_docx_bytes_from_text(text),
                            file_name=f"{os.path.splitext(filename)[0]}_read.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                elif operation == "Write":
                    if not user_text.strip():
                        st.error("Write operation requires non-empty text.")
                    else:
                        new_content = user_text.strip()
                        st.success("Write content ready for download 💾")
                        if download_txt:
                            st.download_button(
                                "Download Written as .txt",
                                data=new_content.encode("utf-8"),
                                file_name=f"{os.path.splitext(filename)[0]}_written.txt",
                                mime="text/plain"
                            )
                        if download_docx and docx is not None:
                            st.download_button(
                                "Download Written as .docx",
                                data=create_docx_bytes_from_text(new_content),
                                file_name=f"{os.path.splitext(filename)[0]}_written.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                elif operation == "Append":
                    if not user_text.strip():
                        st.error("Append operation requires non-empty text.")
                    else:
                        new_content = text + "\n\n" + user_text.strip()
                        st.success("Append content ready for download 💾")
                        if download_txt:
                            st.download_button(
                                "Download Appended as .txt",
                                data=new_content.encode("utf-8"),
                                file_name=f"{os.path.splitext(filename)[0]}_appended.txt",
                                mime="text/plain"
                            )
                        if download_docx and docx is not None:
                            st.download_button(
                                "Download Appended as .docx",
                                data=create_docx_bytes_from_text(new_content),
                                file_name=f"{os.path.splitext(filename)[0]}_appended.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
